import io
import re
from pypdf import PdfReader
from docx import Document

def clean_text(text: str) -> str:
    """
    Cleans raw text by normalizing whitespace, quotes, and removing control chars.
    """
    if not text:
        return ""
    # Normalize spaces and newlines
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[\r\u2028\u2029]', '\n', text)
    # Replace curly quotes and apostrophes
    text = text.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
    
    # Process line by line to keep paragraph structures
    paragraphs = text.split('\n')
    cleaned_paragraphs = []
    for p in paragraphs:
        p_cleaned = re.sub(r'\s+', ' ', p).strip()
        if p_cleaned:
            cleaned_paragraphs.append(p_cleaned)
    return "\n\n".join(cleaned_paragraphs)

def extract_text_from_bytes(file_bytes: bytes, file_name: str) -> str:
    """
    Extracts text from bytes based on the file extension.
    """
    ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
    
    if ext == 'pdf':
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return clean_text("\n".join(text_parts))
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
            
    elif ext == 'docx':
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            text_parts = [para.text for para in doc.paragraphs if para.text]
            # Handle tables in word documents if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return clean_text("\n".join(text_parts))
        except Exception as e:
            raise ValueError(f"Failed to parse Word Document: {str(e)}")
            
    else:
        # Default to raw text (txt)
        for encoding in ('utf-8', 'latin-1', 'utf-16'):
            try:
                return clean_text(file_bytes.decode(encoding))
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with standard encodings.")